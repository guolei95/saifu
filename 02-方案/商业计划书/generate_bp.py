"""
赛赋（SaiFu）商业计划书 - Word文档生成脚本
基于13章框架 + 官方模板格式规范
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── 工具函数 ───

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def add_placeholder(doc, text="⚠️ 此处内容待填写"):
    """添加占位符段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.italic = True
    run.font.size = Pt(10)
    return p

def add_normal_para(doc, text, bold=False, size=11, space_after=6):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    return h

def add_info_table_row(table, label, value):
    """信息表添加一行"""
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    # 设置第一列宽度
    row.cells[0].width = Cm(3)
    # 样式
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    return row

# ─── 主程序 ───

def create_business_plan():
    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ── 默认字体设置 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ── 标题样式 ──
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '黑体'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            h_style.font.size = Pt(16)
        elif i == 2:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    # ===========================================
    # 封面页
    # ===========================================
    for _ in range(6):
        doc.add_paragraph()

    # 大赛名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('中国国际大学生创新大赛')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    doc.add_paragraph()

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—' * 30)
    run.font.size = Pt(10)

    doc.add_paragraph()

    # 商业计划书
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('商 业 计 划 书')
    run.font.size = Pt(28)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235)  # 品牌蓝

    doc.add_paragraph()

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—' * 30)
    run.font.size = Pt(10)

    for _ in range(4):
        doc.add_paragraph()

    # 项目信息表
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    info_data = [
        ('项目名称', '赛赋（SaiFu）—— AI驱动的竞赛全流程辅助平台'),
        ('团队名称', '⚠️ 待填写'),
        ('负责人', '郭雷'),
        ('所在省份', '湖北省'),
        ('所属高校', '湖北师范大学文理学院'),
        ('联系方式', ''),
    ]
    for label, value in info_data:
        add_info_table_row(table, label, value)

    # 让封面表居中
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2025年  ·  创意组')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 分页：目录 ──
    add_page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    doc.add_paragraph()

    # TOC 提示
    p = doc.add_paragraph()
    run = p.add_run('（提示：在 Word 中按 Ctrl+A 全选 → 右键 → 更新域，即可自动生成目录）')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # TOC 域
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z </w:instrText>')
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run('（请在此处右键 → 更新域，生成目录）')
    run4.font.color.rgb = RGBColor(150, 150, 150)
    run4.font.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    # ===========================================
    # 第1章 执行摘要
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '一、执行摘要', level=1)
    add_normal_para(doc, '⚠️ 最后写 —— 等所有章节完成后浓缩。此处为框架占位。', bold=True, size=10)

    add_heading_styled(doc, '1.1 项目名称与定位', level=2)
    add_placeholder(doc, '⚠️ 赛赋是什么，一句话说清。')

    add_heading_styled(doc, '1.2 核心价值主张', level=2)
    add_placeholder(doc, '⚠️ 解决什么问题，怎么解决。')

    add_heading_styled(doc, '1.3 市场机会', level=2)
    add_placeholder(doc, '⚠️ 双非3000万+在校生，竞赛参与率低。')

    add_heading_styled(doc, '1.4 产品概览', level=2)
    add_placeholder(doc, '⚠️ 三方向 + 双路径入口。')

    add_heading_styled(doc, '1.5 商业模式摘要', level=2)
    add_placeholder(doc, '⚠️ C端免费增值 + B端高校采购。')

    add_heading_styled(doc, '1.6 团队亮点', level=2)
    add_placeholder(doc, '⚠️ 双非学生最懂双非痛点。')

    # ===========================================
    # 第2章 项目背景与需求分析
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '二、项目背景与需求分析', level=1)
    add_normal_para(doc, '负责人：商业队友A  |  字数目标：1500-2000字  |  评审权重：⭐⭐⭐（30分-创新维度）', bold=False, size=9)

    add_heading_styled(doc, '2.1 宏观背景', level=2)
    add_placeholder(doc, '⚠️ 互联网+大赛扩容、双创教育政策、竞赛与保研挂钩。')

    add_heading_styled(doc, '2.2 目标用户画像', level=2)
    add_placeholder(doc, '⚠️ 双非大一到大三学生，4个典型画像。')

    add_heading_styled(doc, '2.3 四大核心痛点', level=2)
    add_placeholder(doc, '⚠️ 信息差、分不清、不匹配、组队难。')

    add_heading_styled(doc, '2.4 现有方案缺陷', level=2)
    add_placeholder(doc, '⚠️ 赛氪10年只做信息聚合，无AI辅助。')

    add_heading_styled(doc, '2.5 需求验证', level=2)
    add_placeholder(doc, '⚠️ 问卷数据 + 访谈原话。')

    add_heading_styled(doc, '2.6 需求优先级矩阵', level=2)
    add_placeholder(doc, '⚠️ 选赛推荐 > 方案写作 > 备赛规划 > 组队匹配。')

    # ===========================================
    # 第3章 市场分析
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '三、市场分析', level=1)
    add_normal_para(doc, '负责人：商业队友A  |  字数目标：1000-1200字  |  评审权重：⭐⭐（15分-商业维度）', bold=False, size=9)

    add_heading_styled(doc, '3.1 市场规模测算', level=2)
    add_placeholder(doc, '⚠️ 高校数量、双非占比、目标用户规模。')

    add_heading_styled(doc, '3.2 黄石政策红利', level=2)
    add_placeholder(doc, '⚠️ 数字化转型试点 + AI城市计划 + 创业补贴。')

    # ===========================================
    # 第4章 产品设计与功能
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '四、产品设计与功能', level=1)
    add_normal_para(doc, '负责人：小雷  |  字数目标：1200-1500字  |  评审权重：⭐⭐⭐（30分-创新维度）', bold=False, size=9)

    add_heading_styled(doc, '4.1 产品架构总览', level=2)
    add_placeholder(doc, '⚠️ 三方向（大创/商赛/数模）+ 通用层。配图：产品架构总览图（请设计队友出图）。')

    add_heading_styled(doc, '4.2 用户路径设计', level=2)
    add_placeholder(doc, '⚠️ 路径A「我知道要参加什么」/ 路径B「我不知道」。配图：用户路径流程图（请设计队友出图）。')

    add_heading_styled(doc, '4.3 核心功能', level=2)
    add_placeholder(doc, '⚠️ 选赛推荐、方案生成、赛后沉淀（3个核心详写，其余一笔带过）。')

    add_heading_styled(doc, '4.4 产品形态', level=2)
    add_placeholder(doc, '⚠️ Web应用 + 微信小程序（MVP先做网页版）。')

    add_heading_styled(doc, '4.5 核心性能指标', level=2)
    add_placeholder(doc, '⚠️ 准确率、响应时间、覆盖量等量化数据。【小雷建Word表格在此】')

    # ===========================================
    # 第5章 技术方案与壁垒
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '五、技术方案与壁垒', level=1)
    add_normal_para(doc, '负责人：小雷  |  字数目标：1000-1200字  |  评审权重：⭐⭐⭐（30分-创新维度）', bold=False, size=9)

    add_heading_styled(doc, '5.1 技术架构图', level=2)
    add_placeholder(doc, '⚠️ 前端→后端API→AI引擎层→知识库/提示词库。配图：技术架构图（请设计队友出图）。')

    add_heading_styled(doc, '5.2 核心AI能力', level=2)
    add_placeholder(doc, '⚠️ 大语言模型 + RAG检索增强 + 结构化提示词体系。')

    add_heading_styled(doc, '5.3 知识库构建', level=2)
    add_placeholder(doc, '⚠️ 赛事库(300+赛事) + 方案模板库 + 往届案例库。')

    add_heading_styled(doc, '5.4 技术壁垒分析', level=2)
    add_placeholder(doc, '⚠️ 提示词壁垒 + 知识库壁垒 + 数据飞轮效应。')

    add_heading_styled(doc, '5.5 已有技术成果', level=2)
    add_placeholder(doc, '⚠️ V1 CLI运行截图 + 提示词体系规模。')

    # ===========================================
    # 第6章 商业模式
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '六、商业模式', level=1)
    add_normal_para(doc, '负责人：商业队友B  |  字数目标：800-1000字  |  评审权重：⭐⭐（15分-商业维度）', bold=False, size=9)

    add_heading_styled(doc, '6.1 商业模式画布', level=2)
    add_placeholder(doc, '⚠️ 客户细分 + 价值主张 + 渠道 + 收入来源。配图：商业模式画布（请设计队友出图，标准9格画布）。')

    add_heading_styled(doc, '6.2 用户分层与收费', level=2)
    add_placeholder(doc, '⚠️ C端免费 / C端付费 / B端高校采购。')

    add_heading_styled(doc, '6.3 收入模型', level=2)
    add_placeholder(doc, '⚠️ 4条收入线：C端订阅 + B端合作 + 企业合作 + 增值服务。')

    add_heading_styled(doc, '6.4 付费逻辑', level=2)
    add_placeholder(doc, '⚠️ 为什么用户愿意付费？对标ChatGPT Plus，定价19.9元/月（学生心理价位）。')

    # ===========================================
    # 第7章 营销策略与竞品分析
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '七、营销策略与竞品分析', level=1)
    add_normal_para(doc, '负责人：商业队友B  |  字数目标：800-1000字  |  评审权重：⭐⭐（15分-商业维度）', bold=False, size=9)

    add_heading_styled(doc, '7.1 SWOT分析', level=2)
    add_placeholder(doc, '⚠️ 内部优劣势 + 外部机会威胁。')

    add_heading_styled(doc, '7.2 竞品分析', level=2)
    add_placeholder(doc, '⚠️ 赛氪 vs 保研岛 vs 晨星专利 vs 赛赋。【商业B建Word表格：4竞品×6维度】')

    add_heading_styled(doc, '7.3 获客路径', level=2)
    add_placeholder(doc, '⚠️ 种子用户（校内社群）→ 双非联盟 → KOL合作。配图：获客路径图（请设计队友出图）。')

    add_heading_styled(doc, '7.4 分阶段推广计划', level=2)
    add_placeholder(doc, '⚠️ 试点→扩张→规模化。')

    # ===========================================
    # 第8章 教育实效
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '八、教育实效', level=1)
    add_normal_para(doc, '负责人：小雷（主）+ 商业队友A/B（各100字）  |  字数目标：800-900字  |  评审权重：⭐⭐⭐（30分-教育维度）', bold=False, size=9)

    add_heading_styled(doc, '8.1 个人成长与专创融合', level=2)
    add_placeholder(doc, '⚠️ 小雷写500字（能力成长线+专创融合表）+ 商业A/B各写100字（我在项目中学到了什么）。【小雷建Word表格：4专业×3列专创融合表】')

    add_heading_styled(doc, '8.2 示范引领', level=2)
    add_placeholder(doc, '⚠️ 经验分享计划，带动更多同学参赛。')

    add_heading_styled(doc, '8.3 学校支持', level=2)
    add_placeholder(doc, '⚠️ 双创学院/孵化基地对接情况。')

    # ===========================================
    # 第9章 社会价值
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '九、社会价值', level=1)
    add_normal_para(doc, '负责人：商业队友A  |  字数目标：800-1000字  |  评审权重：⭐⭐（10分-社会价值）', bold=False, size=9)

    add_heading_styled(doc, '9.1 带动就业', level=2)
    add_placeholder(doc, '⚠️ 直接就业 + 间接就业。')

    add_heading_styled(doc, '9.2 教育公平', level=2)
    add_placeholder(doc, '⚠️ 降低双非学生竞赛参与门槛。')

    add_heading_styled(doc, '9.3 社会效益', level=2)
    add_placeholder(doc, '⚠️ 高等教育普惠、创新创业能力提升。配图：社会价值示意图（请设计队友出图）。')

    # ===========================================
    # 第10章 实施计划与里程碑
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '十、实施计划与里程碑', level=1)
    add_normal_para(doc, '负责人：商业队友B  |  字数目标：1000-1200字  |  评审权重：⭐⭐（15分-商业维度）', bold=False, size=9)

    add_heading_styled(doc, '10.1 MVP阶段', level=2)
    add_placeholder(doc, '⚠️ 调研验证 + 网页MVP搭建 + 软著申请。')

    add_heading_styled(doc, '10.2 迭代阶段', level=2)
    add_placeholder(doc, '⚠️ 方案设计模块完整版 + 知识库100+赛事 + 用户内测。')

    add_heading_styled(doc, '10.3 完善阶段', level=2)
    add_placeholder(doc, '⚠️ 选赛推荐 + 商赛/数模方向 + PPT模块上线。')

    add_heading_styled(doc, '10.4 甘特图/路线图', level=2)
    add_placeholder(doc, '⚠️ 配图。【商业B用Excel自带甘特图模板制作，截图放入】')

    add_heading_styled(doc, '10.5 已完成里程碑', level=2)
    add_normal_para(doc, 'V1 CLI版本 ✅ | 提示词体系24文件 ✅ | 初步PPT ✅')

    add_heading_styled(doc, '10.6 用户验证计划', level=2)
    add_placeholder(doc, '⚠️ 已完成验证 + 进行中内测 + 规划中反馈收集。配图：用户验证三阶段图（请设计队友出图）。')

    add_heading_styled(doc, '10.7 版本迭代与运营', level=2)
    add_placeholder(doc, '⚠️ 反馈收集→改进闭环 + 内容维护。')

    # ===========================================
    # 第11章 财务分析
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '十一、财务分析', level=1)
    add_normal_para(doc, '负责人：商业队友B  |  字数目标：800-1000字  |  评审权重：⭐（15分-商业维度）', bold=False, size=9)

    add_heading_styled(doc, '11.1 初期投入', level=2)
    add_placeholder(doc, '⚠️ 大模型API + 服务器/域名 + 软著申请。')

    add_heading_styled(doc, '11.2 月度运营成本', level=2)
    add_placeholder(doc, '⚠️ API调用费 + 服务器 + 维护。')

    add_heading_styled(doc, '11.3 收入预测', level=2)
    add_placeholder(doc, '⚠️ 6/12/18个月三档预测。【商业B用Excel做好截图放入】')

    add_heading_styled(doc, '11.4 资金来源', level=2)
    add_placeholder(doc, '⚠️ 创业补贴5000 + 省级扶持2-10万 + 自筹。')

    add_heading_styled(doc, '11.5 盈亏平衡分析', level=2)
    add_placeholder(doc, '⚠️ 需要多少付费用户覆盖成本。')

    add_heading_styled(doc, '11.6 融资规划', level=2)
    add_placeholder(doc, '⚠️ 股权结构 + 种子轮/天使轮/A轮计划 + 资金用途。配图：融资路线图（请设计队友出图）。')

    add_heading_styled(doc, '11.7 财务风险控制', level=2)
    add_placeholder(doc, '⚠️ API成本管控 + 付费转化率底线。')

    # ===========================================
    # 第12章 团队介绍
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '十二、团队介绍', level=1)
    add_normal_para(doc, '负责人：小雷（12.1-12.2）+ 商业B（12.3-12.5）  |  字数目标：600-800字', bold=False, size=9)

    add_heading_styled(doc, '12.1 核心成员介绍', level=2)
    add_placeholder(doc, '⚠️ 小雷写：照片+专业+分工+一句话亮点。【小雷排版Word图文混排】')

    add_heading_styled(doc, '12.2 指导老师', level=2)
    add_placeholder(doc, '⚠️ 小雷写：姓名+职称+研究方向+对项目的贡献。')

    add_heading_styled(doc, '12.3 团队优势', level=2)
    add_placeholder(doc, '⚠️ 商业B写：双非视角 + 调研先行 + 技术壁垒清晰。')

    add_heading_styled(doc, '12.4 招募计划', level=2)
    add_placeholder(doc, '⚠️ 商业B写：待补充角色的招募时间表。')

    add_heading_styled(doc, '12.5 专家顾问', level=2)
    add_placeholder(doc, '⚠️ 商业B写：拟邀请的顾问方向、对象、贡献。【商业B建Word表格：3方向×3列】')

    # ===========================================
    # 第13章 风险分析与应对
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '十三、风险分析与应对', level=1)
    add_normal_para(doc, '负责人：商业队友B  |  字数目标：500-800字', bold=False, size=9)

    add_heading_styled(doc, '13.1 技术风险', level=2)
    add_placeholder(doc, '⚠️ API成本高 → 免费模型+提示词优化。')

    add_heading_styled(doc, '13.2 竞争风险', level=2)
    add_placeholder(doc, '⚠️ 赛氪跟进AI → 垂直场景壁垒+数据飞轮。')

    add_heading_styled(doc, '13.3 执行风险', level=2)
    add_placeholder(doc, '⚠️ 团队不全 → MVP先行+同步招募。')

    add_heading_styled(doc, '13.4 合规风险', level=2)
    add_placeholder(doc, '⚠️ AI使用规范 → 主动提醒用户标注。')

    add_heading_styled(doc, '13.5 商业化风险', level=2)
    add_placeholder(doc, '⚠️ 付费意愿待验证 → MVP免费验证留存。')

    # ===========================================
    # 附录
    # ===========================================
    add_page_break(doc)
    add_heading_styled(doc, '附录', level=1)

    add_heading_styled(doc, '附录A  用户调研问卷及结果', level=2)
    add_placeholder(doc, '⚠️ 负责人：商业队友A  |  优先级：⭐⭐')

    add_heading_styled(doc, '附录B  产品原型与界面展示', level=2)
    add_placeholder(doc, '⚠️ 负责人：设计队友  |  优先级：⭐⭐⭐  |  3-5页+设计说明')

    add_heading_styled(doc, '附录C  技术实现证明', level=2)
    add_placeholder(doc, '⚠️ 负责人：小雷  |  优先级：⭐⭐  |  代码截图+架构图')

    add_heading_styled(doc, '附录D  知识产权', level=2)
    add_placeholder(doc, '⚠️ 负责人：小雷  |  优先级：⭐⭐⭐  |  软著受理通知书')

    add_heading_styled(doc, '附录E  参考文献', level=2)
    add_placeholder(doc, '⚠️ 各章负责人汇总。')

    add_heading_styled(doc, '附录F  证明材料', level=2)
    add_placeholder(doc, '⚠️ 全队汇总  |  优先级：⭐⭐')

    # ── 保存 ──
    output_path = '02-方案/商业计划书/0508-赛赋商业计划书-框架版.docx'
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    path = create_business_plan()
    print(f'✅ 商业计划书已生成：{path}')
    print(f'   文件大小：{os.path.getsize(path) / 1024:.1f} KB')
