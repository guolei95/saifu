"""
Markdown → Word (.docx) 转换器
处理标题、表格、代码块、复选框、引用块、粗体斜体等
"""
import re
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_borders(table):
    """给表格添加边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_horizontal_rule(doc):
    """添加水平分割线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    # 减小段落间距
    pSpacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="60" w:after="60"/>'
    )
    pPr.append(pSpacing)

def add_run_with_formatting(paragraph, text):
    """解析行内格式并添加 run"""
    # 处理 **bold** *italic* `code` 混合格式
    # 使用正则拆分
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?))'
    parts = re.findall(pattern, text)

    for full, bold_text, italic_text, code_text, plain in parts:
        if plain and plain.strip():
            # 去掉可能有 trailing 的内容
            run = paragraph.add_run(plain)
        elif full.startswith('**') and bold_text:
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif full.startswith('*') and not full.startswith('**') and italic_text:
            run = paragraph.add_run(italic_text)
            run.italic = True
        elif full.startswith('`') and code_text:
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            # 灰色背景用 highlighting
            run.font.highlight_color = None  # 暂时不用高亮

def is_table_separator(line):
    """判断是否是表格分隔行 |---|---|"""
    return bool(re.match(r'^\|[\s\-:|]+\|$', line))

def parse_table_row(line):
    """解析表格行，返回单元格列表"""
    # 去掉首尾的 |
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    # 按 | 分割，但保留转义
    cells = [c.strip() for c in line.split('|')]
    return cells

def create_table_from_data(doc, headers, rows):
    """根据表头和行数据创建表格"""
    ncols = len(headers)
    nrows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'

    # 填充表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        # 表头背景色
        set_cell_shading(cell, '2B579A')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                # 解析行内格式
                add_run_with_formatting(p, cell_text)
                for run_obj in p.runs:
                    run_obj.font.size = Pt(10)
                    run_obj.font.name = '微软雅黑'
                # 奇偶行交替背景
                if i % 2 == 0:
                    set_cell_shading(cell, 'F2F6FC')

    # 设置列宽
    col_width = Cm(16.5 / ncols)  # A4可用宽度
    for j in range(ncols):
        for row in table.rows:
            row.cells[j].width = col_width

    doc.add_paragraph()  # 表后空行
    return table

def setup_styles(doc):
    """设置文档样式"""
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # 标题样式
    for level, (size, color_hex) in enumerate([
        (22, '1A3C6D'),  # Heading 1
        (16, '2B579A'),  # Heading 2
        (13, '3A6DB5'),  # Heading 3
    ], 1):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = '微软雅黑'
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        h_style.paragraph_format.space_before = Pt(12 if level > 1 else 18)
        h_style.paragraph_format.space_after = Pt(6)

def convert_md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 Word 文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    setup_styles(doc)

    i = 0
    in_code_block = False
    code_lines = []
    in_quote_block = False
    quote_lines = []
    ordered_list_num = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.2
                    run = p.add_run(cl)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                doc.add_paragraph()  # 空行
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 跳过分隔线内联的 --- （非独立行）
        # 先检查是否是独立的分隔线
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # 空行处理
        if not line.strip():
            # 检查下一行是否是新表格的开始
            if i + 1 < len(lines) and lines[i+1].strip().startswith('|') and not is_table_separator(lines[i+1].strip()):
                # 可能是表格开始，继续
                pass
            ordered_list_num = 0
            i += 1
            continue

        stripped = line.strip()

        # 处理表格（需要多行读取）
        if stripped.startswith('|') and i + 1 < len(lines) and is_table_separator(lines[i+1].strip()):
            # 这是一个表格
            headers = parse_table_row(stripped)
            # 跳过分隔行
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = parse_table_row(lines[i].strip())
                if row:
                    # 确保行长度与表头一致
                    while len(row) < len(headers):
                        row.append('')
                    rows.append(row[:len(headers)])
                i += 1
            create_table_from_data(doc, headers, rows)
            continue

        # 处理引用块
        if stripped.startswith('> '):
            quote_lines.append(stripped[2:])
            i += 1
            # 检查下一行是否还是引用
            if i < len(lines) and lines[i].strip().startswith('>'):
                continue
            # 输出累积的引用块
            for ql in quote_lines:
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(ql)
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            quote_lines = []
            continue

        # 处理标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            p = doc.add_heading(title, level=level)
            ordered_list_num = 0
            i += 1
            continue

        # 处理无序列表（含复选框 - [ ] - [x]）
        ul_match = re.match(r'^-\s+\[([ x])\]\s+(.+)$', stripped)
        if ul_match:
            checked = ul_match.group(1) == 'x'
            text = ul_match.group(2)
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            # 前缀
            prefix = '☑ ' if checked else '☐ '
            run = p.add_run(prefix + text)
            run.font.size = Pt(10.5)
            run.font.name = '微软雅黑'
            ordered_list_num = 0
            i += 1
            continue

        ul_match2 = re.match(r'^-\s+(.+)$', stripped)
        if ul_match2:
            text = ul_match2.group(1)
            p = doc.add_paragraph(text, style='List Bullet')
            ordered_list_num = 0
            i += 1
            continue

        # 处理有序列表
        ol_match = re.match(r'^(\d+)[\.\)]\s+(.+)$', stripped)
        if ol_match:
            ordered_list_num += 1
            text = ol_match.group(2)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 处理特殊标记行（如「> 附：」但已经处理过引用）
        # 处理 🆕 ⭐ 等特殊符号 - 直接当普通段落

        # 普通段落
        p = doc.add_paragraph()
        add_run_with_formatting(p, stripped)
        ordered_list_num = 0
        i += 1

    # 保存
    doc.save(docx_path)
    print(f'[OK] Generated: {docx_path}')


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    files = [
        (r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-商业队友A.md',
         r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-商业队友A.docx'),
        (r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-商业队友B.md',
         r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-商业队友B.docx'),
        (r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-设计队友.md',
         r'D:\我的竞赛项目-AI赋能竞赛系统\02-方案\商业计划书\0507-任务要求-设计队友.docx'),
    ]

    for md_path, docx_path in files:
        if os.path.exists(md_path):
            convert_md_to_docx(md_path, docx_path)
        else:
            print(f'[WARN] File not found: {md_path}')

    print('\nAll done!')
