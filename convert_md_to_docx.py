"""
将三个任务要求 MD 文件转为格式规整的 Word 文档
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = r"D:\我的竞赛项目-AI赋能竞赛系统"
FILES = [
    ("0507-任务要求-商业队友A.md", "0507-任务要求-商业队友A.docx"),
    ("0507-任务要求-商业队友B.md", "0507-任务要求-商业队友B.docx"),
    ("0507-任务要求-设计队友.md", "0507-任务要求-设计队友.docx"),
]
SRC_DIR = os.path.join(BASE, "conversations", "conv-1778215237719")
OUT_DIR = os.path.join(BASE, "02-方案")

# ============== 样式常量 ==============
FONT_NAME = "微软雅黑"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)
COLOR_PRIMARY = RGBColor(0x25, 0x63, 0xEB)  # 品牌蓝
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_TABLE_HEADER = RGBColor(0x25, 0x63, 0xEB)
COLOR_CHECK = RGBColor(0x22, 0xC5, 0x5E)
COLOR_STAR = RGBColor(0xF5, 0x9E, 0x0B)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:color="{val.get("color", "D1D5DB")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def make_table_borders(table):
    """给整个表格加统一边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'<w:left w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'<w:right w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="D1D5DB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False,
                             font_size=None, color=None, alignment=None,
                             space_after=None, space_before=None, font_name=None):
    """添加格式化段落"""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = font_size or FONT_SIZE_BODY
        run.font.color.rgb = color or COLOR_BODY
        run.font.name = font_name or FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT_NAME)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = space_after
    if space_before is not None:
        pf.space_before = space_before
    return p


def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3}
    colors_map = {1: COLOR_PRIMARY, 2: COLOR_PRIMARY, 3: COLOR_HEADING}
    space_before = {1: Pt(24), 2: Pt(20), 3: Pt(14)}
    space_after = {1: Pt(8), 2: Pt(6), 3: Pt(4)}

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = sizes.get(level, FONT_SIZE_BODY)
    run.font.color.rgb = colors_map.get(level, COLOR_HEADING)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    pf = p.paragraph_format
    pf.space_before = space_before.get(level, Pt(8))
    pf.space_after = space_after.get(level, Pt(4))
    if level == 1:
        # 加下划线装饰
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:color="2563EB" w:space="4"/>'
            f'</w:pBdr>'
        )
        p._p.get_or_add_pPr().append(pBdr)
    return p


def add_blockquote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    # 左边框
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:color="2563EB" w:space="8"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = FONT_SIZE_BODY
    run.font.color.rgb = COLOR_MUTED
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return p


def add_code_block(doc, text):
    """添加代码块"""
    for i, line in enumerate(text.strip().split('\n')):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        if i == 0:
            p.paragraph_format.space_before = Pt(4)
        if i == len(text.strip().split('\n')) - 1:
            p.paragraph_format.space_after = Pt(4)
        # 背景色通过段落底纹
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p


def add_hr(doc):
    """添加水平分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:color="E5E7EB" w:space="1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_checkbox_item(doc, text, checked=False):
    """添加复选框项"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    symbol = "☑" if checked else "☐"
    run = p.add_run(f"{symbol}  {text}")
    run.font.size = FONT_SIZE_BODY
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if checked:
        run.font.color.rgb = COLOR_CHECK
    else:
        run.font.color.rgb = COLOR_BODY
    return p


def add_table_from_rows(doc, rows, header=True):
    """从二维数组创建格式化表格"""
    if not rows:
        return None
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= max_cols:
                break
            cell = row.cells[j]
            # 清除默认段落
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

            if header and i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "2563EB")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run.font.color.rgb = COLOR_BODY
                if i % 2 == 1:
                    set_cell_shading(cell, "F9FAFB")

    make_table_borders(table)

    # 设置列宽
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph()  # 表后空行
    return table


def parse_inline_formatting(paragraph, text):
    """解析行内格式 (bold, italic, code, links)"""
    # 处理 **bold**
    parts = []
    remaining = text

    while remaining:
        # 找各种格式标记
        bold_match = re.match(r'\*\*(.+?)\*\*', remaining)
        italic_match = re.match(r'\*(.+?)\*', remaining)
        code_match = re.match(r'`(.+?)`', remaining)
        link_match = re.match(r'\[(.+?)\]\((.+?)\)', remaining)

        if bold_match:
            parts.append(('bold', bold_match.group(1)))
            remaining = remaining[bold_match.end():]
        elif code_match and not bold_match:
            parts.append(('code', code_match.group(1)))
            remaining = remaining[code_match.end():]
        elif link_match:
            parts.append(('link', link_match.group(1)))
            remaining = remaining[link_match.end():]
        elif italic_match:
            parts.append(('italic', italic_match.group(1)))
            remaining = remaining[italic_match.end():]
        else:
            # 普通文本，找到下一个特殊字符
            next_special = re.search(r'\*\*|`|\[|\*', remaining)
            if next_special:
                end = next_special.start()
                if end > 0:
                    parts.append(('normal', remaining[:end]))
                remaining = remaining[end:]
            else:
                parts.append(('normal', remaining))
                break

    for ptype, text in parts:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.size = FONT_SIZE_BODY
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

        if ptype == 'bold':
            run.bold = True
        elif ptype == 'italic':
            run.italic = True
        elif ptype == 'code':
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif ptype == 'link':
            run.font.color.rgb = COLOR_PRIMARY
            run.underline = True
        else:
            run.font.color.rgb = COLOR_BODY


def add_body_paragraph(doc, text):
    """添加正文段落，支持行内格式"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    parse_inline_formatting(p, text)
    return p


def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    table_mode = False
    table_rows = []
    code_mode = False
    code_lines = []
    blockquote_mode = False
    blockquote_lines = []
    list_mode = False

    while i < len(lines):
        line = lines[i]

        # 空行
        if line.strip() == '' or line.strip() == '&nbsp;':
            if table_mode:
                if table_rows:
                    add_table_from_rows(doc, table_rows)
                    table_rows = []
                table_mode = False
            if code_mode:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                code_mode = False
            if blockquote_mode:
                add_blockquote(doc, ' '.join(blockquote_lines))
                blockquote_lines = []
                blockquote_mode = False
            if list_mode:
                if i + 1 < len(lines) and not lines[i + 1].strip().startswith(('- ', '  - ', '    - ')):
                    list_mode = False
            i += 1
            continue

        stripped = line.strip()

        # 代码块开始/结束
        if stripped.startswith('```'):
            if code_mode:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                code_mode = False
            else:
                if table_mode and table_rows:
                    add_table_from_rows(doc, table_rows)
                    table_rows = []
                    table_mode = False
                if blockquote_mode:
                    add_blockquote(doc, ' '.join(blockquote_lines))
                    blockquote_lines = []
                    blockquote_mode = False
                code_mode = True
            i += 1
            continue

        if code_mode:
            if stripped:  # skip language identifier
                code_lines.append(stripped if stripped != '```' else '')
            i += 1
            continue

        # 水平线
        if stripped == '---':
            if table_mode and table_rows:
                add_table_from_rows(doc, table_rows)
                table_rows = []
                table_mode = False
            if blockquote_mode:
                add_blockquote(doc, ' '.join(blockquote_lines))
                blockquote_lines = []
                blockquote_mode = False
            add_hr(doc)
            i += 1
            continue

        # 引用块
        if stripped.startswith('> '):
            blockquote_mode = True
            blockquote_lines.append(stripped[2:])
            # peek next
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                continue
            else:
                add_blockquote(doc, ' '.join(blockquote_lines))
                blockquote_lines = []
                blockquote_mode = False
                i += 1
                continue

        # 表格行
        if stripped.startswith('|') and stripped.endswith('|'):
            if stripped.replace('|', '').replace('-', '').replace(' ', '').strip() == '':
                # 这是分隔行，跳过
                table_mode = True
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            table_mode = True
            i += 1
            continue

        # 如果之前在表格模式但现在不是表格行了
        if table_mode:
            if table_rows:
                add_table_from_rows(doc, table_rows)
                table_rows = []
            table_mode = False

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            # 去除标题中的 emoji/标记
            title_text = re.sub(r'^[^\w一-鿿]+', '', title_text).strip()
            add_heading_styled(doc, title_text, level)
            i += 1
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*]\s+(.+)', stripped)
        if list_match:
            indent_level = len(list_match.group(1))
            list_text = list_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            # bullet
            bullet_run = p.add_run("•  ")
            bullet_run.font.size = FONT_SIZE_BODY
            bullet_run.font.name = FONT_NAME
            bullet_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            parse_inline_formatting(p, list_text)
            list_mode = True
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+[\.\、]\s*(.+)', stripped)
        if ol_match:
            indent_level = len(ol_match.group(1))
            list_text = ol_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            parse_inline_formatting(p, list_text)
            list_mode = True
            i += 1
            continue

        # 复选框
        cb_match = re.match(r'^-\s*\[([ x])\]\s+(.+)', stripped)
        if cb_match:
            checked = cb_match.group(1).lower() == 'x'
            cb_text = cb_match.group(2)
            add_checkbox_item(doc, cb_text, checked)
            i += 1
            continue

        # 普通文本
        if stripped:
            add_body_paragraph(doc, stripped)
        i += 1

    # 处理末尾残留
    if table_mode and table_rows:
        add_table_from_rows(doc, table_rows)
    if code_mode and code_lines:
        add_code_block(doc, '\n'.join(code_lines))
    if blockquote_mode and blockquote_lines:
        add_blockquote(doc, ' '.join(blockquote_lines))

    doc.save(docx_path)
    print(f"✅ 已生成: {docx_path}")


if __name__ == "__main__":
    for md_name, docx_name in FILES:
        md_path = os.path.join(SRC_DIR, md_name)
        docx_path = os.path.join(OUT_DIR, docx_name)
        if not os.path.exists(md_path):
            print(f"⚠️ 文件不存在: {md_path}")
            continue
        convert_md_to_docx(md_path, docx_path)
    print("\n全部完成！")
